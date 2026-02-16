from __future__ import annotations

import json
import logging
import zipfile
from io import BytesIO
from typing import Any

logger = logging.getLogger(__name__)


def _build_tree(nodes: list[dict]) -> dict | None:
    """Build a tree from a flat node list. Returns the root node with 'children' populated."""
    logger.debug("Building tree from %d nodes", len(nodes))
    node_map: dict[str, dict] = {}
    root = None
    for n in nodes:
        node_map[n["id"]] = {**n, "children": []}
    for n in nodes:
        if n["parent_id"] is None:
            root = node_map[n["id"]]
        else:
            parent = node_map.get(n["parent_id"])
            if parent:
                parent["children"].append(node_map[n["id"]])
            else:
                logger.warning("Orphan node: id=%s, parent_id=%s not found", n["id"], n["parent_id"])
    # Sort children by position
    if root:
        _sort_children(root)
        logger.debug("Tree built successfully, root='%s'", root.get("content", ""))
    else:
        logger.warning("No root node found (no node with parent_id=NULL)")
    return root


def _sort_children(node: dict) -> None:
    node["children"].sort(key=lambda c: c.get("position", 0))
    for child in node["children"]:
        _sort_children(child)


def export_docx(map_name: str, nodes: list[dict]) -> BytesIO:
    logger.info("Generating DOCX: map_name='%s', nodes=%d", map_name, len(nodes))
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        logger.error("python-docx is not installed. Run: pip install python-docx")
        raise

    doc = Document()
    root = _build_tree(nodes)
    if not root:
        logger.warning("No root node, creating empty document with title only")
        doc.add_heading(map_name, level=1)
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    doc.add_heading(root["content"] or map_name, level=1)

    def _add_nodes(node: dict, level: int) -> None:
        for child in node["children"]:
            content = child["content"] or ""
            if level <= 8:
                # Heading levels 2-9
                doc.add_heading(content, level=level)
            else:
                # For deep nesting, use indented paragraphs
                indent = level - 9
                p = doc.add_paragraph(content)
                p.paragraph_format.left_indent = Inches(0.3 * indent)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
            _add_nodes(child, level + 1)

    _add_nodes(root, 2)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    logger.info("DOCX generated: %d bytes", buf.getbuffer().nbytes)
    return buf


def export_xlsx(map_name: str, nodes: list[dict]) -> BytesIO:
    logger.info("Generating XLSX: map_name='%s', nodes=%d", map_name, len(nodes))
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment
    except ImportError:
        logger.error("openpyxl is not installed. Run: pip install openpyxl")
        raise

    wb = Workbook()
    ws = wb.active
    ws.title = map_name[:31]  # Excel sheet name max 31 chars

    # Headers
    ws.append(["Level", "Content", "Parent Content"])
    for cell in ws[1]:
        cell.font = Font(bold=True)

    root = _build_tree(nodes)
    if not root:
        logger.warning("No root node, creating empty spreadsheet with headers only")
        buf = BytesIO()
        wb.save(buf)
        buf.seek(0)
        return buf

    # Build parent content lookup
    parent_map: dict[str | None, str] = {}
    for n in nodes:
        parent_map[n["id"]] = n["content"] or ""

    row_count = 0

    def _add_rows(node: dict, level: int, parent_content: str) -> None:
        nonlocal row_count
        content = node["content"] or ""
        indent = "  " * level
        ws.append([level, f"{indent}{content}", parent_content if level > 0 else ""])
        row_count += 1
        for child in node["children"]:
            _add_rows(child, level + 1, content)

    _add_rows(root, 0, "")

    # Auto-adjust column widths
    ws.column_dimensions["A"].width = 8
    ws.column_dimensions["B"].width = 50
    ws.column_dimensions["C"].width = 30

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    logger.info("XLSX generated: %d bytes, %d rows", buf.getbuffer().nbytes, row_count)
    return buf


def export_xmind(map_name: str, nodes: list[dict]) -> BytesIO:
    """Export as XMind 8+ format (.xmind is a ZIP containing content.json and metadata.json)."""
    logger.info("Generating XMind: map_name='%s', nodes=%d", map_name, len(nodes))
    root = _build_tree(nodes)

    def _build_topic(node: dict) -> dict[str, Any]:
        topic: dict[str, Any] = {
            "id": node["id"],
            "title": node["content"] or "",
        }
        if node["children"]:
            topic["children"] = {
                "attached": [_build_topic(c) for c in node["children"]]
            }
        return topic

    root_topic = _build_topic(root) if root else {"id": "root", "title": map_name}

    content = [
        {
            "id": "sheet-1",
            "title": map_name,
            "rootTopic": root_topic,
        }
    ]

    metadata = {
        "creator": {
            "name": "MindMap Export",
            "version": "1.0.0",
        }
    }

    buf = BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("content.json", json.dumps(content, ensure_ascii=False, indent=2))
        zf.writestr("metadata.json", json.dumps(metadata, ensure_ascii=False, indent=2))
    buf.seek(0)
    logger.info("XMind generated: %d bytes", buf.getbuffer().nbytes)
    return buf
